"""Membuat draft laporan proyek UAS dari artefak evaluasi aktual."""

from __future__ import annotations

import csv
import json
import sys
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "deliverables" / "Laporan_Proyek_Chatbot_FAQ_UMB.docx"
EVAL_DIR = ROOT / "artifacts" / "evaluation"
ASSET_DIR = ROOT / "artifacts" / "report_assets"
DOC_SKILL = Path(
    "/Users/koala/.codex/plugins/cache/openai-primary-runtime/"
    "documents/26.723.12215/skills/documents"
)
sys.path.insert(0, str(DOC_SKILL / "scripts"))
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


# standard_business_brief dengan override bernama "A4 academic submission".
NAVY = RGBColor(7, 31, 61)
BLUE = RGBColor(11, 77, 162)
GOLD = RGBColor(202, 139, 20)
INK = RGBColor(25, 39, 58)
MUTED = RGBColor(91, 106, 122)
WHITE = RGBColor(255, 255, 255)
LIGHT_BLUE = "EAF2FC"
LIGHT_GRAY = "F2F4F7"
LINE = "CBD6E3"
BODY_FONT = "Arial"


def set_run_font(run, size=10, bold=False, italic=False, color=INK, name=BODY_FONT):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=LINE, size=5):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), color)


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=8.5, color=MUTED)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    section.header_distance = Cm(0.65)
    section.footer_distance = Cm(0.65)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal.font.size = Pt(9.7)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    heading_tokens = {
        "Heading 1": (14.5, BLUE, 10, 4),
        "Heading 2": (11.5, NAVY, 7, 3),
        "Heading 3": (10.2, NAVY, 5, 2),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = BODY_FONT
        style.font.size = Pt(9.5)
        style.paragraph_format.left_indent = Inches(0.32)
        style.paragraph_format.first_line_indent = Inches(-0.16)
        style.paragraph_format.space_after = Pt(2.5)
        style.paragraph_format.line_spacing = 1.08

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("UAS NLP  |  CHATBOT FAQ AKADEMIK UMB")
    set_run_font(run, size=7.5, bold=True, color=MUTED)
    header.paragraph_format.space_after = Pt(0)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Laporan Proyek  |  ")
    set_run_font(run, size=8.5, color=MUTED)
    add_field(footer, "PAGE")


def add_title_block(doc: Document):
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(5)
    kicker.paragraph_format.space_after = Pt(6)
    run = kicker.add_run("PROJECT AKHIR · NATURAL LANGUAGE PROCESSING")
    set_run_font(run, size=8.5, bold=True, color=GOLD)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("SIAKAD Assist")
    set_run_font(run, size=25, bold=True, color=NAVY, name="Georgia")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run(
        "Chatbot FAQ Akademik Berbasis TF-IDF dan Logistic Regression\n"
        "Universitas Mercu Buana"
    )
    set_run_font(run, size=13, bold=True, color=BLUE)

    metadata = [
        ("Mata kuliah", "Natural Language Processing"),
        ("Kelompok", "Anggota 1 · Anggota 2 · Anggota 3"),
        ("Domain", "FAQ akademik kampus"),
        ("Tanggal", "27 Juli 2026"),
    ]
    table = doc.add_table(rows=1, cols=2)
    for index, (label, value) in enumerate(metadata):
        row = table.rows[0] if index == 0 else table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
    style_table(table, header=False, first_col=True)
    apply_geometry(doc, table, [1.15, 5.75])

    callout = doc.add_paragraph()
    callout.paragraph_format.space_before = Pt(7)
    callout.paragraph_format.space_after = Pt(7)
    callout.paragraph_format.left_indent = Inches(0.12)
    callout.paragraph_format.right_indent = Inches(0.12)
    run = callout.add_run(
        "Ringkasan hasil: 250 utterance seimbang, 5 intent, accuracy 98,00%, "
        "macro F1 97,99%, serta dukungan CLI dan UI Streamlit dengan log CSV."
    )
    set_run_font(run, size=10, bold=True, color=NAVY)
    p_pr = callout._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_BLUE)
    p_pr.append(shd)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_body(doc, text, bold_lead=None):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = paragraph.add_run(text[len(bold_lead) :])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_run_font(run, size=9.5)
    return paragraph


def style_table(table, header=True, first_col=False):
    table.style = "Table Grid"
    if header:
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        repeat = OxmlElement("w:tblHeader")
        repeat.set(qn("w:val"), "true")
        tr_pr.append(repeat)
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            if header and row_index == 0:
                shade_cell(cell, "0B4DA2")
            elif first_col and col_index == 0:
                shade_cell(cell, LIGHT_GRAY)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        size=8.5,
                        bold=(header and row_index == 0) or (first_col and col_index == 0),
                        color=WHITE if header and row_index == 0 else INK,
                    )


def apply_geometry(doc, table, weights):
    section = doc.sections[0]
    width = int(section.page_width.twips - section.left_margin.twips - section.right_margin.twips)
    widths = column_widths_from_weights(weights, width)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=width,
        indent_dxa=110,
        cell_margins_dxa={"top": 70, "bottom": 70, "start": 110, "end": 110},
    )


def add_table(doc, headers, rows, weights):
    table = doc.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = str(header)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
    style_table(table)
    apply_geometry(doc, table, weights)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(1)
    return table


def draw_box(draw, box, label, fill, font, text_color="white"):
    draw.rounded_rectangle(box, radius=18, fill=fill, outline="#cbd6e3", width=3)
    left, top, right, bottom = box
    lines = label.split("\n")
    line_h = font.size + 8
    start_y = (top + bottom - line_h * len(lines)) / 2
    for index, line in enumerate(lines):
        bounds = draw.textbbox((0, 0), line, font=font)
        x = (left + right - (bounds[2] - bounds[0])) / 2
        draw.text((x, start_y + index * line_h), line, fill=text_color, font=font)


def draw_arrow(draw, start, end, color="#0b4da2", width=6):
    draw.line([start, end], fill=color, width=width)
    x, y = end
    draw.polygon([(x, y), (x - 16, y - 10), (x - 16, y + 10)], fill=color)


def build_diagrams():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    font_path = "/System/Library/Fonts/Supplemental/Arial.ttf"
    bold_path = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"
    font = ImageFont.truetype(font_path, 30)
    small = ImageFont.truetype(font_path, 24)
    flow_font = ImageFont.truetype(font_path, 20)
    bold = ImageFont.truetype(bold_path, 30)

    architecture = Image.new("RGB", (1600, 620), "#f3f6fa")
    draw = ImageDraw.Draw(architecture)
    draw.text((60, 35), "Arsitektur pemrosesan chatbot", fill="#071f3d", font=bold)
    boxes = [
        ((55, 155, 255, 290), "Input\npengguna", "#071f3d"),
        ((330, 155, 555, 290), "Preprocessing", "#0b4da2"),
        ((630, 85, 880, 220), "TF-IDF +\nLogistic Regression", "#0b4da2"),
        ((630, 280, 880, 415), "Regex slot +\nstate manager", "#ca8b14"),
        ((955, 155, 1180, 290), "Respons\nchatbot", "#071f3d"),
        ((1290, 85, 1515, 220), "CLI", "#0b4da2"),
        ((1290, 280, 1515, 415), "UI Streamlit", "#0b4da2"),
    ]
    for box, label, fill in boxes:
        draw_box(draw, box, label, fill, small)
    draw_arrow(draw, (255, 222), (330, 222))
    draw_arrow(draw, (555, 222), (630, 152))
    draw_arrow(draw, (555, 222), (630, 347))
    draw_arrow(draw, (880, 152), (955, 222))
    draw_arrow(draw, (880, 347), (955, 222))
    draw_arrow(draw, (1180, 222), (1290, 152))
    draw_arrow(draw, (1180, 222), (1290, 347))
    draw.text((610, 510), "Setiap interaksi -> log CSV tersamarkan", fill="#5d6a78", font=small)
    architecture.save(ASSET_DIR / "architecture.png", quality=95)

    flow = Image.new("RGB", (1600, 520), "white")
    draw = ImageDraw.Draw(flow)
    draw.text((60, 28), "Flow dialog pengisian KRS", fill="#071f3d", font=bold)
    flow_boxes = [
        ((50, 175, 285, 315), "IDLE\nDeteksi intent", "#071f3d"),
        ((365, 175, 620, 315), "WAITING_FOR_NIM\nValidasi regex", "#0b4da2"),
        ((700, 175, 980, 315), "SELECTING_MATKUL\nEkstraksi nomor", "#0b4da2"),
        ((1060, 175, 1350, 315), "WAITING_CONFIRMATION\nYa / batal", "#ca8b14"),
        ((1415, 175, 1565, 315), "IDLE", "#071f3d"),
    ]
    for box, label, fill in flow_boxes:
        draw_box(draw, box, label, fill, flow_font)
    for start, end in [
        ((285, 245), (365, 245)),
        ((620, 245), (700, 245)),
        ((980, 245), (1060, 245)),
        ((1350, 245), (1415, 245)),
    ]:
        draw_arrow(draw, start, end)
    draw.text((1200, 355), "Input ambigu -> ulangi konfirmasi", fill="#5d6a78", font=small)
    flow.save(ASSET_DIR / "dialog_flow.png", quality=95)


def add_figure(doc, path, caption, width_inches=6.55):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(1)
    cap.paragraph_format.space_after = Pt(5)
    cap.paragraph_format.keep_with_next = True
    run = cap.add_run(caption)
    set_run_font(run, size=8.2, italic=True, color=MUTED)


def load_csv(name):
    with (EVAL_DIR / name).open(encoding="utf-8", newline="") as handle:
        return list(csv.DictReader(handle))


def build_report():
    build_diagrams()
    summary = json.loads((EVAL_DIR / "evaluation_summary.json").read_text(encoding="utf-8"))
    distribution = load_csv("dataset_distribution.csv")
    examples = load_csv("preprocessing_examples.csv")
    mistakes = load_csv("misclassified_examples.csv")

    doc = Document()
    configure_document(doc)
    add_title_block(doc)

    add_heading(doc, "1. Pendahuluan dan Identifikasi Masalah", 1)
    add_body(
        doc,
        "Mahasiswa membutuhkan jawaban cepat untuk pertanyaan berulang tentang KRS, "
        "pembayaran UKT, jadwal ujian, beasiswa, dan akses portal. Informasi tersebar "
        "pada beberapa kanal sehingga pertanyaan sederhana tetap membebani layanan "
        "akademik. Proyek ini merancang chatbot NLP sederhana yang memberi jawaban awal "
        "dan memandu simulasi KRS tanpa menggantikan keputusan petugas kampus."
    )
    add_heading(doc, "1.1 Tujuan dan Batas Sistem", 2)
    add_body(
        doc,
        "Tujuan sistem adalah mengenali intent pertanyaan mahasiswa, mengekstrak slot "
        "NIM serta pilihan mata kuliah, mempertahankan konteks percakapan, meminta "
        "konfirmasi, dan menyimpan log. Sistem menggunakan data dummy dan tidak melakukan "
        "transaksi atau perubahan pada SIAKAD produksi."
    )
    add_heading(doc, "1.2 Kebutuhan Fungsional", 2)
    functional_rows = [
        ("F-01", "Menerima teks dari CLI dan UI Streamlit", "Terpenuhi"),
        ("F-02", "Mengklasifikasikan lima intent FAQ", "Terpenuhi"),
        ("F-03", "Mengekstrak NIM dan nomor mata kuliah dengan regex", "Terpenuhi"),
        ("F-04", "Menjalankan dialog KRS multi-turn dan konfirmasi", "Terpenuhi"),
        ("F-05", "Memberi fallback untuk confidence rendah/input invalid", "Terpenuhi"),
        ("F-06", "Menyimpan log CSV dengan masking NIM", "Terpenuhi"),
    ]
    add_table(doc, ["ID", "Kebutuhan", "Status"], functional_rows, [0.75, 5.15, 1.0])

    add_heading(doc, "2. Teori Singkat NLP", 1)
    add_body(
        doc,
        "Natural Language Processing (NLP) mempelajari metode komputasional untuk "
        "memproses bahasa manusia. Pada chatbot berbasis intent, teks dinormalisasi lalu "
        "diubah menjadi fitur numerik. TF-IDF memberi bobot tinggi pada istilah yang "
        "informatif dalam sebuah dokumen tetapi relatif jarang pada keseluruhan korpus "
        "[1][2]."
    )
    add_body(
        doc,
        "Logistic Regression memperkirakan probabilitas kelas dan sesuai untuk fitur teks "
        "sparse. Prediksi dievaluasi melalui accuracy, precision, recall, F1-score, dan "
        "confusion matrix. Precision menilai ketepatan prediksi positif, recall menilai "
        "cakupan data kelas yang ditemukan, sedangkan F1 menyeimbangkan keduanya [3][4]."
    )

    dataset_heading = add_heading(doc, "3. Dataset dan Preprocessing", 1)
    dataset_heading.paragraph_format.page_break_before = True
    add_body(
        doc,
        f"Dataset berisi {summary['dataset_size']} utterance berbahasa Indonesia. Lima intent "
        "dibuat seimbang untuk mengurangi bias kelas. Variasi mencakup bahasa formal, "
        "percakapan mahasiswa, singkatan, typo ringan, dan susunan pertanyaan berbeda."
    )
    add_table(
        doc,
        ["Intent", "Jumlah", "Proporsi"],
        [
            (row["intent"], row["count"], f"{int(row['count']) / summary['dataset_size']:.0%}")
            for row in distribution
        ],
        [4.7, 1.1, 1.1],
    )
    add_figure(
        doc,
        EVAL_DIR / "dataset_distribution.png",
        "Gambar 1. Distribusi dataset per intent (masing-masing 50 utterance).",
        width_inches=5.9,
    )
    add_heading(doc, "3.1 Tahap Preprocessing", 2)
    for text in [
        "Lowercase: menyamakan kapitalisasi.",
        "Cleaning: menghapus tanda baca dan karakter khusus.",
        "Tokenization: memecah teks berdasarkan token kata.",
        "Normalisasi: memperbaiki singkatan/typo seperti bgmn menjadi bagaimana.",
    ]:
        add_bullet(doc, text)
    selected_examples = [examples[index] for index in [0, 2, 4, 6, 8]]
    add_table(
        doc,
        ["Intent", "Sebelum", "Sesudah"],
        [(row["intent"], row["text"], row["clean_text"]) for row in selected_examples],
        [1.7, 2.6, 2.6],
    )

    add_heading(doc, "4. Desain dan Implementasi", 1)
    add_figure(doc, ASSET_DIR / "architecture.png", "Gambar 2. Arsitektur pemrosesan chatbot.")
    add_body(
        doc,
        "Representasi teks menggabungkan word n-gram (1-2) dan character n-gram (2-5) "
        "melalui FeatureUnion. Character n-gram membuat model lebih tahan terhadap typo, "
        "sedangkan word n-gram mempertahankan pola makna. Logistic Regression menggunakan "
        "class_weight=balanced, C=2,0, dan random_state=42."
    )
    add_heading(doc, "4.1 Slot Filling dan Dialog Manager", 2)
    add_body(
        doc,
        "Slot filling menggunakan regex untuk mencari NIM 10-12 digit serta nomor mata "
        "kuliah. DialogManager menyimpan state dan konteks sehingga pesan berikutnya dapat "
        "ditafsirkan sesuai tahap proses. Jawaban konfirmasi yang ambigu tidak langsung "
        "membatalkan proses, tetapi meminta pengguna menjawab ya atau batal."
    )
    add_figure(doc, ASSET_DIR / "dialog_flow.png", "Gambar 3. State diagram alur KRS multi-turn.")
    add_heading(doc, "4.2 Struktur Data Percakapan dan Log", 2)
    add_body(
        doc,
        "Dataset intent menggunakan pasangan {text, intent}. Database demo memuat mahasiswa, "
        "program studi, semester, katalog mata kuliah, SKS, dan biaya. Log CSV menyimpan "
        "timestamp, session ID, kanal, pesan pengguna, respons, intent, confidence, state "
        "sebelum, dan state sesudah. NIM disamarkan sebelum log ditulis."
    )
    add_heading(doc, "4.3 Implementasi CLI dan UI", 2)
    add_body(
        doc,
        "CLI dijalankan dengan python3 cli_app.py dan menyediakan bantuan, reset, keluar, "
        "serta metadata prediksi. UI Streamlit memakai elemen chat resmi, quick actions, "
        "indikator state, unduh log, dan tampilan responsif tanpa aset eksternal [5]."
    )

    add_heading(doc, "5. Evaluasi dan Analisis", 1)
    add_body(
        doc,
        f"Dataset dibagi secara stratified menjadi {summary['train_size']} data latih dan "
        f"{summary['test_size']} data uji (80:20). Seed 42 digunakan agar hasil dapat "
        "direproduksi. Tabel berikut memakai macro average agar setiap intent memiliki "
        "bobot yang sama."
    )
    metric_rows = [
        ("Accuracy", f"{summary['accuracy'] * 100:.2f}%"),
        ("Precision (macro)", f"{summary['macro_precision'] * 100:.2f}%"),
        ("Recall (macro)", f"{summary['macro_recall'] * 100:.2f}%"),
        ("F1-Score (macro)", f"{summary['macro_f1'] * 100:.2f}%"),
        ("F1-Score (weighted)", f"{summary['weighted_f1'] * 100:.2f}%"),
    ]
    add_table(doc, ["Metrik", "Hasil"], metric_rows, [5.2, 1.7])
    add_figure(
        doc,
        EVAL_DIR / "confusion_matrix.png",
        "Gambar 4. Confusion matrix pada 50 data uji.",
        width_inches=5.75,
    )
    add_heading(doc, "5.1 Intent yang Paling Sering Salah", 2)
    if mistakes:
        mistake = mistakes[0]
        add_body(
            doc,
            f"Terdapat {len(mistakes)} kesalahan. Intent {mistake['actual_intent']} salah "
            f"diprediksi sebagai {mistake['predicted_intent']} pada kalimat \"{mistake['text']}\". "
            "Kata 'daftar' dan frasa waktu 'semester depan' juga sering muncul pada konteks "
            "pendaftaran beasiswa, sehingga fitur leksikal kedua kelas beririsan."
        )
    else:
        add_body(doc, "Tidak terdapat salah klasifikasi pada split pengujian ini.")
    add_heading(doc, "5.2 Penyebab Kesalahan dan Keterbatasan", 2)
    for text in [
        "Dataset masih disusun untuk cakupan proyek dan belum berasal dari log produksi nyata.",
        "Model mengandalkan kemiripan leksikal sehingga pertanyaan sangat pendek atau ambigu dapat tertukar.",
        "Sistem hanya mendukung lima intent serta dua jenis slot utama.",
        "Confidence probabilistik belum dikalibrasi khusus dan threshold 0,35 bersifat empiris.",
        "Informasi biaya, jadwal, syarat, dan portal harus diverifikasi melalui kanal resmi kampus.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "6. Pengujian Sistem", 1)
    add_body(
        doc,
        "Automated tests mencakup validitas dataset, preprocessing, pipeline model, slot "
        "filling, state machine, konfirmasi ambigu, masking NIM, logger, CLI, dan UI. "
        "Pengujian manual menggunakan tiga skenario: FAQ UKT, jadwal ujian dengan follow-up "
        "portal, serta alur KRS lengkap sampai konfirmasi."
    )
    add_table(
        doc,
        ["Skenario", "Input utama", "Hasil yang diharapkan"],
        [
            ("FAQ UKT", "cara bayar ukt melalui bca", "Intent pembayaran_ukt dan respons prosedur"),
            ("Follow-up", "jadwal uas ... / portalnya di mana", "Konteks diarahkan ke akses_portal"),
            ("KRS multi-turn", "KRS / NIM / 1,2 / ya", "State selesai dan konfirmasi berhasil"),
        ],
        [1.3, 2.5, 3.1],
    )

    add_heading(doc, "7. Kesimpulan", 1)
    add_body(
        doc,
        "Chatbot FAQ Akademik berhasil memenuhi komponen utama rubrik: dataset 250 data, "
        "preprocessing, TF-IDF, Logistic Regression, regex slot filling, dialog multi-turn "
        "dengan konfirmasi, evaluasi lengkap, CLI, UI, dan log CSV. Accuracy 98,00% "
        "menunjukkan model efektif untuk dataset proyek, tetapi perlu perluasan data nyata "
        "dan integrasi sumber resmi sebelum diterapkan sebagai layanan kampus."
    )

    add_heading(doc, "8. Pembagian Tugas", 1)
    add_table(
        doc,
        ["Anggota", "Peran", "Kontribusi"],
        [
            ("Anggota 1", "Data & NLP", "Dataset, preprocessing, pelatihan model"),
            ("Anggota 2", "Backend", "Slot filling, dialog manager, CLI, logging"),
            ("Anggota 3", "UI & Evaluasi", "Streamlit, pengujian, laporan, slide"),
        ],
        [1.2, 1.7, 4.0],
    )

    add_heading(doc, "Daftar Pustaka", 1)
    references = [
        "[1] D. Jurafsky dan J. H. Martin, Speech and Language Processing, edisi ke-3, draft daring, Stanford University, 2026. https://web.stanford.edu/~jurafsky/slp3/ed3book.pdf",
        "[2] Scikit-learn Developers, TfidfVectorizer Documentation. https://scikit-learn.org/stable/modules/generated/sklearn.feature_extraction.text.TfidfVectorizer.html",
        "[3] Scikit-learn Developers, LogisticRegression Documentation. https://scikit-learn.org/stable/modules/generated/sklearn.linear_model.LogisticRegression.html",
        "[4] Scikit-learn Developers, Metrics and Scoring: Quantifying the Quality of Predictions. https://scikit-learn.org/stable/modules/model_evaluation.html",
        "[5] Streamlit, Chat Elements API Documentation. https://docs.streamlit.io/develop/api-reference/chat",
    ]
    for reference in references:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.18)
        paragraph.paragraph_format.first_line_indent = Inches(-0.18)
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(reference)
        set_run_font(run, size=8.3, color=MUTED)

    doc.core_properties.title = "Laporan Proyek Chatbot FAQ Akademik UMB"
    doc.core_properties.subject = "Project Akhir Natural Language Processing"
    doc.core_properties.author = "Kelompok UAS NLP - Universitas Mercu Buana"
    doc.core_properties.keywords = "NLP, chatbot, TF-IDF, Logistic Regression, Streamlit"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"DOCX dibuat: {OUTPUT}")


if __name__ == "__main__":
    build_report()
